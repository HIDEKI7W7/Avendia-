import re
import asyncio
from io import BytesIO
from docx import Document as DocxDocument
import google.generativeai as genai
from sqlmodel.ext.asyncio.session import AsyncSession
from src.chatbot.search import retrieve_relevant_context

def add_formatted_text_to_paragraph(paragraph, text: str) -> None:
    """
    Parsea fragmentos de texto en negrita (**texto**) y los añade
    como runs enriquecidos al párrafo de Word de manera limpia.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Remover los asteriscos de Markdown
            clean_text = part[2:-2]
            run = paragraph.add_run(clean_text)
            run.bold = True
        else:
            paragraph.add_run(part)

def parse_markdown_to_docx(markdown_content: str, doc: DocxDocument) -> None:
    """
    Parsea línea por línea el contenido Markdown y crea elementos nativos
    de Word (Títulos, Listas, Párrafos con negritas).
    """
    lines = markdown_content.split('\n')
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        # Comprobar Títulos (# Título)
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', line_str)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2)
            # Normalizar nivel de título al rango soportado de Word (1 a 4)
            doc.add_heading(title_text, level=min(level, 4))
            continue
            
        # Comprobar listas (* Elemento)
        list_match = re.match(r'^[\*\-\+]\s+(.*)$', line_str)
        if list_match:
            item_text = list_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text_to_paragraph(p, item_text)
            continue
            
        # Párrafo normal
        p = doc.add_paragraph()
        add_formatted_text_to_paragraph(p, line_str)

async def generate_document_content(
    user_id: str,
    doc_type: str,
    form_data: dict,
    session: AsyncSession
) -> str:
    """
    Estructura la solicitud para Gemini inyectando el contexto de las normativas
    de la base de datos RAG y los datos del formulario docente/directivo.
    """
    # Generar query para consultar el RAG basándonos en los campos del formulario
    query_context = f"Plantilla de {doc_type} escolar con datos " + ", ".join(
        f"{k}: {v}" for k, v in form_data.items()
    )
    
    # Obtener el contexto normativo
    context, _ = await retrieve_relevant_context(
        query_text=query_context,
        session=session,
        limit=2
    )
    
    # Armar prompt estructurado
    prompt = (
        f"Eres un asistente administrativo y pedagógico experto. Escribe el cuerpo de un documento formal de tipo '{doc_type}' "
        f"utilizando la siguiente información:\n\n"
        f"[DATOS DEL FORMULARIO DOCENTE/DIRECTIVO]:\n"
        + "\n".join(f"- {k.replace('_', ' ').capitalize()}: {v}" for k, v in form_data.items())
        + f"\n\n[CONTEXTO NORMATIVO ESCOLAR]:\n{context}\n\n"
        f"REQUISITOS DE FORMATO (CRÍTICO):\n"
        f"1. Redacta el documento completo en formato Markdown utilizando únicamente etiquetas estándar:\n"
        f"   - Títulos y secciones con '#' o '##'.\n"
        f"   - Párrafos limpios separados por saltos de línea.\n"
        f"   - Enfatiza palabras clave o datos usando negritas '**'.\n"
        f"   - Listas o acuerdos usando viñetas '*'.\n"
        f"2. Usa un lenguaje profesional y estructurado.\n"
        f"3. No incluyas comentarios iniciales ni explicaciones adicionales, devuelve únicamente el texto del documento.\n"
    )
    
    # Invocar Gemini
    model = genai.GenerativeModel("gemini-3.5-flash")
    loop = asyncio.get_running_loop()
    
    response = await loop.run_in_executor(
        None,
        lambda: model.generate_content(prompt)
    )
    
    # Capturar consumo de tokens
    prompt_tokens = 0
    completion_tokens = 0
    if hasattr(response, "usage_metadata") and response.usage_metadata:
        prompt_tokens = response.usage_metadata.prompt_token_count
        completion_tokens = response.usage_metadata.candidates_token_count
    else:
        # Fallback de estimación básica (1 token aprox 4 caracteres)
        prompt_tokens = len(prompt) // 4
        completion_tokens = len(response.text or "") // 4
        
    total_tokens = prompt_tokens + completion_tokens
    
    # Calcular costo en USD y Soles (tasa 3.8)
    cost_usd = (prompt_tokens * 0.000000075) + (completion_tokens * 0.00000030)
    cost_soles = cost_usd * 3.80
    
    from src.models.ai_log import AILog
    from sqlalchemy import text
    import uuid
    
    ai_log = AILog(
        user_id=uuid.UUID(user_id),
        model_name="gemini-3.5-flash",
        prompt_tokens=prompt_tokens,
        completion_tokens=completion_tokens,
        total_tokens=total_tokens,
        cost_usd=cost_usd,
        cost_soles=cost_soles
    )
    session.add(ai_log)
    
    # Descontar créditos al usuario (ej: 10 créditos por documento generado)
    await session.execute(
        text("UPDATE users SET credits = GREATEST(0, credits - 10) WHERE id = :id"),
        {"id": uuid.UUID(user_id)}
    )
    
    # Registrar el acceso más reciente
    from datetime import datetime
    await session.execute(
        text("UPDATE users SET last_access = :now WHERE id = :id"),
        {"now": datetime.utcnow(), "id": uuid.UUID(user_id)}
    )
    
    return response.text or ""


def create_editable_docx(title: str, markdown_content: str) -> BytesIO:
    """
    Inicializa un documento Word en memoria y traduce el markdown generado
    por Gemini a elementos de estilos nativos de python-docx.
    """
    doc = DocxDocument()
    
    # Cabecera o Título principal
    doc.add_heading(title, 0)
    
    # Parsear y volcar texto enriquecido
    parse_markdown_to_docx(markdown_content, doc)
    
    # Guardar en buffer en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
